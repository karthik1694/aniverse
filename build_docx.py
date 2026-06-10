#!/usr/bin/env python3
"""Generate a simple, easy-to-understand Word (.docx) document explaining
DataStage .ds dataset hand-off and how to export the data, with trusted sources."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "DataStage_Dataset_Guide.docx"

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x40, 0x40, 0x40)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = GREY


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(text, size=15, color=PRIMARY, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run(text)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], "1F4E79")
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if r_idx % 2 == 1:
                shade_cell(cells[i], "DCE6F1")
            para = cells[i].paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = GREY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ---------------------------------------------------------------- Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("DataStage Datasets (.ds) Made Simple")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = PRIMARY

sub = doc.add_paragraph()
sr = sub.add_run("How data moves between jobs, and how to get it out of DataStage")
sr.italic = True
sr.font.size = Pt(11.5)
sr.font.color.rgb = ACCENT

meta = doc.add_paragraph()
mr = meta.add_run("A plain-English guide  |  Prepared by Kiro  |  June 2026")
mr.font.size = Pt(9)
mr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xDC, 0xE6, 0xF1)

# ---------------------------------------------------------------- The big idea
heading("The Big Idea (in one line)")
body("When you have several DataStage jobs running one after another, each job saves its "
     "result so the next job can pick it up. The usual way to save that result is a special "
     "DataStage file called a Dataset, which ends in .ds")

heading("A Simple Picture")
body("Imagine 3 jobs that run in order. Each one hands its data to the next using a .ds file:")
body("Job 1  ->  saves  data1.ds   ->  Job 2  reads it  ->  saves  data2.ds   ->  Job 3  reads it",
     space_after=8)
body("A \u201CSequence Job\u201D is the controller that decides the order: run Job 1, then Job 2, then Job 3.")

# ---------------------------------------------------------------- What is .ds
heading("What Exactly Is a .ds File?")
body("Here is the surprising part: the .ds file is NOT the real data. It is just a small "
     "\u201Cpointer\u201D or \u201Cmap\u201D file. The real data lives in separate hidden files that DataStage "
     "spreads across its machines (nodes) so it can work fast and in parallel.")
body("Think of the .ds file like a library catalog card: the card tells you where the book is, "
     "but the card is not the book itself.")
bullet("The .ds file = the small map (it holds the column info and where the data is).")
bullet("The real data = larger hidden files split across several machines.")
bullet("Important: never copy or delete .ds files using normal commands (like cp or rm). "
       "Use DataStage's own tools, or you will break the link to the real data.")

# ---------------------------------------------------------------- Can you download
heading("Can I Download the .ds File and Open It Elsewhere?")
body("Short answer: No.")
body("The .ds format is private to DataStage. If you copy it to your computer and try to open "
     "it in Excel, Notepad, Python, or anything else, it will not work. Only DataStage can read "
     "its own datasets (using the Data Set stage).")

# ---------------------------------------------------------------- How to get data out
heading("So How Do I Actually Get the Data Out?")
body("You convert the data into a normal file that any tool can open. Here are your options, "
     "from easiest/best to more advanced:")

make_table(
    ["Way to export", "What you get", "Good for"],
    [
        ["Sequential File stage (CSV)",
         "A normal CSV / text file",
         "BEST choice. Open in Excel, Python, anything. Easy to download."],
        ["Database target",
         "A table in a database",
         "When the next step is a database or BI tool."],
        ["orchadmin (command line)",
         "Quick view of the data",
         "Peeking / checking only. Can cut off (truncate) large data."],
        ["Data Set Management (the built-in tool)",
         "View data inside DataStage",
         "Looking at contents without exporting."],
    ],
    col_widths=[2.0, 2.2, 2.6],
)

heading("The Recommended Trick", size=13, color=ACCENT, space_before=6)
body("Let Job 1 save TWO outputs at the same time:")
bullet("A .ds file -> fast hand-off to the next job (keep this for DataStage).")
bullet("A CSV file (Sequential File stage) -> this is the one you download and use anywhere.")
body("So the flow looks like:")
body("Source  ->  Transformer  ->  Data Set (.ds)  [for Job 2]\n"
     "                          ->  Sequential File (.csv)  [for you to download]")

# ---------------------------------------------------------------- Comparison
heading("Quick Compare: DataStage vs. Databricks")
body("If you already know Databricks, here is how the same ideas line up:")
make_table(
    ["Idea", "DataStage", "Databricks"],
    [
        ["The whole pipeline", "Sequence Job", "Workflow"],
        ["One step in it", "Job", "Task (notebook/script)"],
        ["How big data is passed", ".ds Dataset / table / CSV", "Delta / Parquet files"],
        ["Best save format", ".ds (fast, keeps structure)", "Delta (reliable, fast)"],
        ["Passing small values", "Job parameters", "dbutils.jobs.taskValues"],
    ],
    col_widths=[2.0, 2.4, 2.4],
)

# ---------------------------------------------------------------- Remember
heading("Just Remember These 4 Things")
bullet("1. A .ds file is only a map; the real data is hidden and split across machines.")
bullet("2. You cannot open a .ds file outside DataStage.")
bullet("3. To get usable data, save a CSV (Sequential File stage) or load it into a database.")
bullet("4. .ds in DataStage works like Delta/Parquet in Databricks.")

# ---------------------------------------------------------------- Sources
doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xDC, 0xE6, 0xF1)
heading("Trusted Sources")
body("These points were checked against IBM's official documentation and well-known "
     "practitioner sources. (Summarized in our own words.)", space_after=4)

sources = [
    "IBM Docs \u2013 Developing DataStage parallel jobs (data sets are saved data, pointed to by a .ds file): "
    "https://www.ibm.com/docs/en/iis/11.7?topic=qualitystage-developing-parallel-jobs",
    "IBM Docs \u2013 Structure of data sets (data is split into segments and partitions): "
    "https://www.ibm.com/docs/en/iis/11.5?topic=mds-structure-data-sets",
    "IBM Docs \u2013 Data Set stage (read/write datasets inside DataStage): "
    "https://www.ibm.com/docs/en/iis/11.7.0?topic=files-data-sets",
    "IBM Docs \u2013 Sequential file (writes a normal flat file you can use elsewhere): "
    "https://dataplatform.cloud.ibm.com/docs/content/dstage/dsnav/topics/sequential_file_connector.html",
    "IBM Support \u2013 orchadmin dump can truncate output (so it is for peeking, not full export): "
    "https://www.ibm.com/support/pages/datastage-orchadmin-dump-command-dataset-produces-truncated-output-which-when-saved-xml-file-extension-has-incorrect-xml-format",
    "Practitioner blog \u2013 orchadmin command usage (delete, copy, describe, dump): "
    "https://riteshkumargupta.wordpress.com/2012/04/14/processing-infosphere-datastage-dataset-from-command-line-orchadmin-utility/",
    "Practitioner community (Spiceworks) \u2013 datasets are internal format, only readable in DataStage: "
    "https://community.spiceworks.com/t/reading-datastage-datasets/888161",
]
for i, s in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[{i}] {s}")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

doc.save(OUTPUT)
print(f"Word document written to {OUTPUT}")
