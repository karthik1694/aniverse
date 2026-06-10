#!/usr/bin/env python3
"""Generate a clear, step-by-step Word (.docx) guide for exporting an
IN-BETWEEN DataStage job (e.g., Job 2) to a CSV file you can download."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "DataStage_Export_InBetween_Job_to_CSV.docx"

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x40, 0x40, 0x40)
WARN = RGBColor(0xB0, 0x4A, 0x00)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = GREY


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(text, size=15, color=PRIMARY, before=12, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, after=6, color=GREY, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.color.rgb = color
    r.bold = bold
    return p


def bullet(text, color=GREY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.color.rgb = color
    return p


def mono(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = PRIMARY
    return p


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "1F4E79")
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if r_idx % 2 == 1:
                shade(cells[i], "DCE6F1")
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = GREY
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ---------------------------------------------------------- Title
title = doc.add_paragraph()
tr = title.add_run("How to Export an In-Between Job (e.g., Job 2) to CSV in DataStage")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = PRIMARY

sub = doc.add_paragraph()
sr = sub.add_run("Step-by-step guide for getting intermediate data out of a sequence of jobs")
sr.italic = True
sr.font.size = Pt(11.5)
sr.font.color.rgb = ACCENT

meta = doc.add_paragraph()
mr = meta.add_run("Prepared by Kiro  |  June 2026")
mr.font.size = Pt(9)
mr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xDC, 0xE6, 0xF1)

# ---------------------------------------------------------- Intro
heading("What This Guide Covers")
body("In a DataStage sequence of jobs (Job 1 -> Job 2 -> Job 3), you may want the data "
     "produced by an in-between job, such as Job 2, not just the final result. This guide "
     "shows how to export that intermediate data to a CSV file that you can download to your "
     "laptop and open in Excel, Notepad, or Python.")

body("Key idea: you build one small, separate export job that reads Job 2's dataset (.ds) and "
     "writes a CSV. Then you plug that export job into the sequence right after Job 2. You do "
     "NOT need to change Job 1, Job 2, or Job 3.")

# ---------------------------------------------------------- Before you start
heading("Before You Start - One Important Check")
body("Your export job reads the file Job 2 produced (for example job2_output.ds). That file "
     "must actually exist on disk when the export runs.", after=4)
bullet("OK: If Job 2 writes a real, persistent Data Set as its output (the normal pattern), "
       "you are fine.")
bullet("Watch out: If Job 2 only streams data straight into Job 3 in memory (no .ds saved), "
       "there is no file to read. In that case, make Job 2 land a Data Set first.")
p = doc.add_paragraph()
r = p.add_run("Quick check: if you can see job2_output.ds in the Data Set Management tool, "
              "you can export it.")
r.italic = True
r.font.color.rgb = WARN

# ---------------------------------------------------------- Part 1
heading("Part 1 - Build the Export Job")

heading("Step 1: Create a new Parallel Job", size=12.5, color=ACCENT, before=8, after=3)
bullet("Open DataStage Designer.")
bullet("Go to File -> New -> Parallel Job.")
bullet("Save it with a clear name, e.g. ExportJob2CSV.")

heading("Step 2: Add the two stages", size=12.5, color=ACCENT, before=8, after=3)
body("From the palette on the right, drag onto the canvas:", after=3)
bullet("Data Set stage (under the File section) - this READS Job 2's existing .ds")
bullet("Sequential File stage (under the File section) - this WRITES the CSV")

heading("Step 3: Connect them with a link", size=12.5, color=ACCENT, before=8, after=3)
body("Hover over the Data Set stage, then click and drag a link to the Sequential File stage. "
     "You should now see:", after=3)
mono("[ Data Set ]  ----link---->  [ Sequential File ]\n"
     " (job2_output.ds)               (job2_output.csv)")

heading("Step 4: Configure the Data Set stage (the SOURCE)", size=12.5, color=ACCENT, before=8, after=3)
bullet("Double-click the Data Set stage.")
body("In Properties, set File = the path to the dataset produced by Job 2, for example:", after=3)
mono("/datasets/job2_output.ds")
bullet("This is the key change for an in-between export: point it at Job 2's .ds, not the final job's.")
bullet("Go to the Columns tab -> click Load to pull in the column definitions (the schema).")
bullet("Click OK.")

heading("Step 5: Configure the Sequential File stage (the TARGET / CSV)", size=12.5, color=ACCENT, before=8, after=3)
bullet("Double-click the Sequential File stage.")
body("Properties tab -> File = where to write the CSV, for example:", after=3)
mono("/exports/job2_output.csv")
body("Format tab - set these so it is a proper CSV:", after=3)
bullet("Delimiter = comma")
bullet("Quote = double (handles values that contain commas)")
bullet("Final delimiter = end")
bullet("First line is column names = True (gives you a header row)")
body("Advanced tab -> Execution mode = Sequential", after=2, bold=True)
p = doc.add_paragraph()
r = p.add_run("Important: this forces ONE single clean CSV file. If left on Parallel, DataStage "
              "writes one file per node and you get several pieces.")
r.italic = True
r.font.color.rgb = WARN
bullet("Click OK.")

heading("Step 6: Compile the export job", size=12.5, color=ACCENT, before=8, after=3)
bullet("Click the Compile button (the hammer/gear icon).")
bullet("Fix any errors until it says 'Compiled successfully'.")

# ---------------------------------------------------------- Part 2
heading("Part 2 - Add the Export Job After Job 2 in the Sequence")

heading("Step 7: Open your Sequence Job", size=12.5, color=ACCENT, before=8, after=3)
bullet("Open the existing Sequence Job that runs Job 1 -> Job 2 -> Job 3.")

heading("Step 8: Add a Job Activity for the export job", size=12.5, color=ACCENT, before=8, after=3)
bullet("From the palette, drag a Job Activity stage onto the canvas.")
bullet("Double-click it.")
bullet("In Job name, select your new ExportJob2CSV.")
bullet("Click OK.")

heading("Step 9: Connect it after Job 2 - choose one option", size=12.5, color=ACCENT, before=8, after=3)
body("Option A - Insert in the middle (export runs, then Job 3 continues):", after=3, bold=True)
mono("[Job1] -> [Job2] -> [ExportJob2CSV] -> [Job3]")
bullet("Connect Job 2 -> ExportJob2CSV with an 'OK (Conditional)' trigger.")
bullet("Then connect ExportJob2CSV -> Job 3.")
bullet("Effect: Job 3 waits until the CSV is written.")

body("Option B - Branch it off (recommended; does NOT delay Job 3):", after=3, bold=True)
mono("[Job1] -> [Job2] --+--> [Job3]\n"
     "                   +--> [ExportJob2CSV]")
bullet("Connect Job 2 -> Job 3 as normal.")
bullet("ALSO connect Job 2 -> ExportJob2CSV (a second trigger link).")
bullet("Set both triggers to 'OK (Conditional)'.")
bullet("Effect: after Job 2 finishes, Job 3 and the export run together; the export does not "
       "slow the main pipeline.")

heading("Step 10: Compile the Sequence Job", size=12.5, color=ACCENT, before=8, after=3)
bullet("Click Compile on the Sequence.")
bullet("Make sure it compiles cleanly.")

# ---------------------------------------------------------- Part 3
heading("Part 3 - Run and Collect Your CSV")

heading("Step 11: Run the Sequence", size=12.5, color=ACCENT, before=8, after=3)
bullet("Run the Sequence Job (from Designer or Director).")
bullet("It runs Job 1 -> Job 2 -> (export) -> Job 3.")
body("When everything turns green, the CSV is created on the server at:", after=3)
mono("/exports/job2_output.csv")

heading("Step 12: Download the CSV to your laptop", size=12.5, color=ACCENT, before=8, after=3)
body("The file is on the DataStage server, so bring it to your PC using:", after=3)
bullet("WinSCP / FileZilla / MobaXterm -> connect to the server -> drag job2_output.csv to your computer, or")
body("Command line:", after=3)
mono("scp username@datastage-server:/exports/job2_output.csv  C:\\Users\\You\\Downloads\\")

heading("Step 13: Open it", size=12.5, color=ACCENT, before=8, after=3)
bullet("Open the CSV in Excel, Notepad, or Python - done!")

# ---------------------------------------------------------- Summary table
heading("Summary Checklist")
make_table(
    ["Step", "Action"],
    [
        ["1", "New Parallel Job -> name it ExportJob2CSV"],
        ["2", "Add Data Set stage + Sequential File stage"],
        ["3", "Link Data Set -> Sequential File"],
        ["4", "Data Set stage -> point to /datasets/job2_output.ds, Load columns"],
        ["5", "Sequential File -> CSV path, comma delimiter, header = True, Sequential mode"],
        ["6", "Compile the export job"],
        ["7", "Open the Sequence Job"],
        ["8", "Add a Job Activity -> select ExportJob2CSV"],
        ["9", "Connect after Job 2 (Option A inline, or Option B branch) with OK (Conditional)"],
        ["10", "Compile the Sequence"],
        ["11", "Run the Sequence -> CSV created on server"],
        ["12", "Download with WinSCP / scp"],
        ["13", "Open in Excel"],
    ],
    widths=[0.7, 5.8],
)

# ---------------------------------------------------------- Tips
heading("Extra Tip: Export Several Jobs at Once")
body("If you want CSVs of more than one step, just add more export jobs, each pointed at a "
     "different .ds and branched off its job:")
mono("[Job1] -> [Job2] -> [Job3]\n"
     "   |         |         |\n"
     "   v         v         v\n"
     " job1.csv  job2.csv  job3.csv")
body("Each export job is the same small Data Set -> Sequential File job, just pointed at a "
     "different dataset.")

# ---------------------------------------------------------- Sources
doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xDC, 0xE6, 0xF1)
heading("Trusted Sources")
body("Checked against IBM official documentation and well-known practitioner sources. "
     "Summarized in our own words for licensing compliance.", after=4)
sources = [
    "IBM Docs - Data Set stage (read/write datasets inside DataStage): "
    "https://www.ibm.com/docs/en/iis/11.7.0?topic=files-data-sets",
    "IBM Docs - Sequential file (writes a normal flat file; single file when run sequentially): "
    "https://dataplatform.cloud.ibm.com/docs/content/dstage/dsnav/topics/sequential_file_connector.html",
    "IBM Docs - Structure of data sets (data is split into segments and partitions): "
    "https://www.ibm.com/docs/en/iis/11.5?topic=mds-structure-data-sets",
    "IBM Docs - Developing DataStage parallel jobs (datasets are saved data pointed to by a .ds file): "
    "https://www.ibm.com/docs/en/iis/11.7?topic=qualitystage-developing-parallel-jobs",
]
for i, s in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[{i}] {s}")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

doc.save(OUTPUT)
print(f"Word document written to {OUTPUT}")
